"""
PDF + DOCX loaders.

Returns plain text with line breaks preserved — the section splitter relies on
heading lines being isolated. Two-column resumes are merged column-by-column,
not interleaved. For PDFs we try pdfplumber first (better column handling);
pymupdf is the fallback for awkward layouts pdfplumber chokes on.
"""
from __future__ import annotations

import re
from pathlib import Path


# ── PUBLIC ──────────────────────────────────────────────────────────────────

def load_resume(path: str | Path) -> str:
    """Load a resume from disk → normalized plain text. Raises on unsupported format."""
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        text = _load_pdf(p)
    elif suffix in (".docx", ".doc"):
        if suffix == ".doc":
            raise ValueError(
                f"Legacy .doc not supported (requires LibreOffice). Convert to .docx first: {p}"
            )
        text = _load_docx(p)
    elif suffix in (".txt", ".md"):
        text = p.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported format {suffix}: {p}")
    return _normalize_text(text)


# ── PDF ─────────────────────────────────────────────────────────────────────

def _load_pdf(path: Path) -> str:
    """Try column-aware pymupdf extraction first (handles 2-column resume layouts
    where pdfplumber interleaves columns). Fall back to pdfplumber on failure,
    and to plain pymupdf as a final option."""
    try:
        return _load_pdf_columns(path)
    except Exception as e:
        print(f"[loader] column-aware loader failed on {path.name}: {e}; falling back to pdfplumber")
    try:
        return _load_pdf_pdfplumber(path)
    except Exception as e:
        print(f"[loader] pdfplumber failed on {path.name}: {e}; falling back to plain pymupdf")
        return _load_pdf_pymupdf(path)


def _load_pdf_columns(path: Path) -> str:
    """Column-aware extraction using pymupdf's block-level layout info.

    For each page:
      1. Get all text blocks with their (x0, y0, x1, y1) bounding boxes.
      2. Look for a vertical column-separator gap: an x-range that no block
         spans, wider than 5% of page width. If found, blocks left of the gap
         are column A and blocks right of it are column B.
      3. Within each column, sort top-to-bottom and concatenate.
      4. Two columns become two text regions joined by a blank line — the
         section splitter will pick up section headings inside each.

    For single-column pages, blocks are simply read in (y, x) order.
    """
    import fitz

    out_pages: list[str] = []
    doc = fitz.open(path)
    try:
        for page in doc:
            blocks = page.get_text("blocks")
            # block tuple: (x0, y0, x1, y1, text, block_no, block_type)
            # block_type == 0 → text (1 is image)
            text_blocks = [
                b for b in blocks
                if len(b) >= 7 and b[6] == 0 and b[4] and b[4].strip()
            ]
            if not text_blocks:
                continue

            page_width = page.rect.width
            split_x = _find_column_split(text_blocks, page_width)

            if split_x is not None:
                left = sorted(
                    [b for b in text_blocks if (b[0] + b[2]) / 2 < split_x],
                    key=lambda b: b[1],
                )
                right = sorted(
                    [b for b in text_blocks if (b[0] + b[2]) / 2 >= split_x],
                    key=lambda b: b[1],
                )
                page_text = (
                    "\n".join(b[4].strip() for b in left)
                    + "\n\n"
                    + "\n".join(b[4].strip() for b in right)
                )
            else:
                ordered = sorted(text_blocks, key=lambda b: (b[1], b[0]))
                page_text = "\n".join(b[4].strip() for b in ordered)

            out_pages.append(page_text)
    finally:
        doc.close()

    return "\n\n".join(out_pages)


def _find_column_split(blocks: list, page_width: float) -> float | None:
    """Return the x-coordinate of a vertical gap between columns, or None if
    the page is single-column.

    Approach: merge all block x-ranges, find the largest x-gap between
    consecutive merged ranges. If the gap is wider than 5% of page width
    AND positioned near the middle (20-80% of page width), it's a column
    separator."""
    if len(blocks) < 4:
        return None

    ranges = sorted([(b[0], b[2]) for b in blocks])
    merged = [list(ranges[0])]
    for s, e in ranges[1:]:
        if s <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], e)
        else:
            merged.append([s, e])

    if len(merged) < 2:
        return None

    best_gap_width = 0.0
    best_gap_mid = 0.0
    for i in range(1, len(merged)):
        gap_width = merged[i][0] - merged[i - 1][1]
        gap_mid = (merged[i - 1][1] + merged[i][0]) / 2
        # Only consider gaps in the middle 60% of the page — sidebars at
        # the extreme edges are layout decoration, not column structure.
        if 0.2 * page_width < gap_mid < 0.8 * page_width and gap_width > best_gap_width:
            best_gap_width = gap_width
            best_gap_mid = gap_mid

    if best_gap_width < 0.05 * page_width:
        return None
    return best_gap_mid


def _load_pdf_pdfplumber(path: Path) -> str:
    import pdfplumber

    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            # Layout-aware extraction preserves indentation and line breaks
            txt = page.extract_text(layout=False, x_tolerance=2, y_tolerance=3) or ""
            chunks.append(txt)
    return "\n".join(chunks)


def _load_pdf_pymupdf(path: Path) -> str:
    import fitz  # pymupdf

    chunks: list[str] = []
    doc = fitz.open(path)
    try:
        for page in doc:
            chunks.append(page.get_text("text"))
    finally:
        doc.close()
    return "\n".join(chunks)


# ── DOCX ────────────────────────────────────────────────────────────────────

def _load_docx(path: Path) -> str:
    """python-docx → text. Walk paragraphs in order, then tables (resumes often
    use single-cell tables for sidebars). Headings get blank-line padding so
    the section splitter sees them on their own line."""
    from docx import Document

    doc = Document(str(path))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style or "title" in style:
            lines.append("")
            lines.append(text)
            lines.append("")
        else:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(cell_text)

    return "\n".join(lines)


# ── NORMALIZATION ───────────────────────────────────────────────────────────

_BULLETS = "•●▪◦∙·○"  # common bullet glyphs that confuse heading detection
_LIGATURES = {
    "ﬀ": "ff", "ﬁ": "fi", "ﬂ": "fl",
    "ﬃ": "ffi", "ﬄ": "ffl",
}
# Smart-quote / smart-apostrophe normalization — PDFs convert these from Word.
# Untouched, they break regexes that expect straight quotes (Varun's
# "March'24" was actually "March’\u24" in the PDF).
_SMART_PUNCT = {
    "‘": "'", "’": "'",  # ‘ ’
    "‚": "'", "‛": "'",
    "“": '"', "”": '"',  # “ ”
    "„": '"', "‟": '"',
    "′": "'", "″": '"',  # prime marks
}


def _normalize_text(text: str) -> str:
    """Light normalization — collapse repeated blank lines, fix common PDF glyphs,
    strip bullets at line start. Preserves the line-break structure that section
    detection depends on."""
    if not text:
        return ""

    # Fix common PDF ligatures
    for src, dst in _LIGATURES.items():
        text = text.replace(src, dst)

    # Normalize smart quotes / apostrophes to straight (uniform regex matching)
    for src, dst in _SMART_PUNCT.items():
        text = text.replace(src, dst)

    # Strip leading bullet glyphs (with optional whitespace) from each line
    bullet_re = re.compile(rf"^[\s\t]*[{re.escape(_BULLETS)}\-\*]\s*", re.MULTILINE)
    text = bullet_re.sub("", text)

    # Normalize whitespace within lines but keep newlines
    lines = [re.sub(r"[ \t]+", " ", ln).strip() for ln in text.splitlines()]
    text = "\n".join(lines)

    # Collapse 3+ blank lines → 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()
